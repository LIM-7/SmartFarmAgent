# -*- coding: utf-8 -*-
"""通用 Markdown -> Word(.docx) 转换脚本
用法: python md2docx.py <input.md> <output.docx>
支持: #/##/### 标题、表格、无序/有序列表、复选框、引用、代码块、**加粗**、`行内代码`
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, latin="Times New Roman", east="宋体", size=None, bold=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def add_rich_text(paragraph, text, east="宋体", size=10.5):
    """处理 **加粗** 与 `行内代码`"""
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = paragraph.add_run(tok[2:-2])
            set_run_font(r, east=east, size=size, bold=True)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            r = paragraph.add_run(tok[1:-1])
            set_run_font(r, latin="Consolas", east=east, size=size)
        else:
            r = paragraph.add_run(tok)
            set_run_font(r, east=east, size=size)


def add_table(doc, rows):
    """rows: list[list[str]]，首行为表头"""
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, cell_text in enumerate(rows[0]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_rich_text(p, cell_text, east="黑体", size=10)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(ncols):
            val = row[i] if i < len(row) else ""
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_rich_text(p, val, size=10)
    doc.add_paragraph()


def convert(md_path, out_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    # 全局字体
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for hname in ("Heading 1", "Heading 2", "Heading 3"):
        h = doc.styles[hname]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0, 0, 0)
        h._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    in_code = False
    table_buf = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_run_font(r, latin="Consolas", east="宋体", size=9)
            i += 1
            continue

        # 表格
        if stripped.startswith("|"):
            table_buf.append([c.strip() for c in stripped.strip("|").split("|")])
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                cleaned = [r for r in table_buf if not all(re.fullmatch(r"[\s:\-]+", c or "") for c in r)]
                if cleaned:
                    add_table(doc, cleaned)
                table_buf = []
            continue

        # 标题
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:].strip())
            set_run_font(r, east="黑体", size=16, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        m = re.match(r"^(#{2,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1)) - 1
            p = doc.add_heading("", level=level)
            add_rich_text(p, m.group(2).strip(), east="黑体", size=(14 if level == 1 else 12))
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            r = p.add_run(stripped.lstrip("> ").strip())
            set_run_font(r, east="宋体", size=10)
            r.font.italic = True
            i += 1
            continue

        # 复选框
        if stripped.startswith("- [ ]"):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, "□ " + stripped[6:].strip())
            i += 1
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:].strip())
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, m.group(2).strip())
            i += 1
            continue

        # 普通段落
        if stripped:
            p = doc.add_paragraph()
            add_rich_text(p, stripped)
        i += 1

    doc.save(out_path)
    print("OK:", out_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
